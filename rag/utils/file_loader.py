from PyPDF2 import PdfReader
from docx import Document

def extract_text(file):
    if file.name.endswith('.pdf'):
        reader = PdfReader(file)
        #readers.page is a list sort of and 
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)
# return "n/.join([page.extract_text() for page in reader.pages if page.extract_text()])"    

    elif file.name.endswith('.docx'):
        doc = Document(file)
        docs = []
        for para in doc.paragraphs:
            docu= para.text
            if docu:  # If the paragraph is not empty
                docs.append(docu)
        return "\n".join(docs)


        # return "\n".join([para.text for para in doc.paragraphs])
    
  
    elif file.name.endswith('.txt'):
        return file.read().decode('utf-8')
    
    else:
        return ""
